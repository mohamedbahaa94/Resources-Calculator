import base64
from io import BytesIO

import docx
import pandas as pd
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_paragraph_with_bold(doc, text):
    """
    Adds a paragraph to the document with text that can include **bold** segments.
    """
    paragraph = doc.add_paragraph()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if i % 2 == 1:  # Odd indices are the bold segments
            run = paragraph.add_run(part)
            run.bold = True
        else:
            paragraph.add_run(part)
    return paragraph

def add_table_with_no_split(doc, df, title, description=None):
    """
    Adds a table to the document and ensures it does not split across pages.
    Allows an optional description to be included above the table with support for bold text formatting.

    Args:
        doc (Document): The Word document to add the table to.
        df (DataFrame): The data to populate the table.
        title (str): The title for the table.
        description (str, optional): An explanatory paragraph to add before the table. Default is None.
    """
    # Add heading
    doc.add_heading(title, level=3)

    # Add explanatory text if provided
    if description:
        add_paragraph_with_bold(doc, description)

    # Create the table
    table = doc.add_table(rows=1, cols=len(df.columns), style='CustomTableStyle')

    # Add table header
    header_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        header_cells[i].text = column_name

    # Add table rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Prevent the table from splitting across pages
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

    # Ensure the table as a whole does not split
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblPr.append(OxmlElement('w:tblInd'))  # Prevent splitting at the table level

    return table

def generate_introduction(customer_name, high_availability, nas_backup_details, tier_3_disks, gpu_specs):

        """
        Generate a polished narrative-style introduction with refined text and conditional tiering logic
        for PaxeraHealth imaging solutions, using bullet points for key features.
        """
        # Base introduction paragraph
        intro_text = (
            f"The **{customer_name} IT Infrastructure, Hardware, and Virtual Machine (VM) Design Recommendations** document provides a comprehensive blueprint "
            "tailored for the deployment and scalability of **PaxeraHealth’s advanced medical imaging solutions**. "
            "This document outlines the architectural design, hardware specifications, storage strategies, and virtual machine configurations "
            "required to meet the operational and performance needs of healthcare facilities.\n\n"
            "The design prioritizes:\n"
            "• **Scalability:** Ensuring the system can grow seamlessly with increased demand.\n"
            "• **Reliability:** Delivering continuous and uninterrupted services across critical infrastructure.\n"
            "• **Performance Optimization:** Leveraging cutting-edge technology for optimal efficiency.\n"
            "• **Robust Security Measures:** Safeguarding sensitive medical data with industry-standard protocols.\n"
            "• **Compliance:** Adhering to international standards for healthcare IT infrastructure."
        )

        # Initialize a list to hold all sections
        sections = [intro_text]

        # Modular Features Section
        feature_intro_text = "\n**Key Features of the Proposed Design Include:**"
        sections.append(feature_intro_text)

        # High Availability or Single Server
        if high_availability:
            ha_text = (
                "• **High Availability (HA):** Ensures uninterrupted operations, minimizes downtime, and supports fault tolerance, "
                "providing continuous accessibility to critical medical imaging systems. "
                "The design incorporates **automatic failover and failback capabilities**, allowing seamless transitions between primary and backup systems "
                "to maintain service continuity during failures and ensure rapid recovery without manual intervention."
            )
            sections.append(ha_text)
        else:
            single_deployment_text = (
                "• **Single Deployment Architecture:** Designed for environments where redundancy is not a primary requirement. "
                "This deployment typically involves one or two servers operating without redundancy, while still delivering essential system performance "
                "and supporting critical medical imaging workflows efficiently."
            )
            sections.append(single_deployment_text)
        # GPU Specs
        if gpu_specs:
            gpu_text = (
                "• **GPU-Powered Processing:** Enables advanced AI functionalities, optimized performance for resource-intensive tasks, "
                "and innovative AI-driven analysis and visualization."
            )
            sections.append(gpu_text)

        # Storage Logic: Include both Short-Term and Standard Tiering if tier_3_disks is enabled
        if tier_3_disks != 0:
            storage_text = (
                "• **Three-Tier Storage Architecture:** The design incorporates a three-tier storage strategy to optimize imaging workflows and data lifecycle management. "
                "Tier 1 utilizes **high-performance SSD storage** for OS and critical application data, ensuring rapid response times and system stability. "
                "Tier 2 employs **flash storage for short-term imaging workflows**, providing fast access to frequently used imaging data. "
                "Tier 3 leverages **cost-efficient HDD storage** for long-term archiving, ensuring scalability and reliable access to historical imaging data."
            )
        else:
            storage_text = (
                "• **Standard Two-Tier Storage Architecture:** The design features a two-tier storage approach for efficient imaging data management. "
                "Tier 1 utilizes **high-performance SSD storage** for OS and critical application data, ensuring system responsiveness and reliability. "
                "Tier 2 employs **cost-efficient HDD storage** for long-term archiving, offering scalability and secure storage for historical imaging data."
            )
        sections.append(storage_text)
        # NAS Backup
        if nas_backup_details:
            nas_text = (
                "• **NAS Backup Solution:** Provides long-term data protection, robust redundancy, and seamless recovery for critical medical imaging archives."
            )
            sections.append(nas_text)

        # Conclude the introduction
        conclusion_text = (
            f"\nThis document serves as a reference for stakeholders, IT administrators, and decision-makers, providing clarity on system requirements, "
            "deployment guidelines, and long-term operational strategies. Through this approach, **PaxeraHealth** aims to deliver not just a solution "
            f"but a sustainable foundation for future healthcare innovations at **{customer_name}**."
        )
        sections.append(conclusion_text)

        # Join all sections for the introduction
        return sections


# Add Introduction to Document
def add_introduction_to_document(doc, intro_sections):
    """
    Adds the introduction sections to the Word document with bold formatting
    and proper spacing between paragraphs.
    """
    doc.add_heading("Introduction", level=2)
    for section in intro_sections:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.space_after = Pt(12)  # Add space after each paragraph

        # Split the section text by '**' to identify bold segments
        parts = section.split('**')
        for i, part in enumerate(parts):
            run = paragraph.add_run(part)
            if i % 2 == 1:  # Odd indices correspond to bold text
                run.bold = True
def get_binary_file_downloader_html(bin_file, file_label='File', customer_name='Customer', file_extension='.docx'):
    file_name = f'{customer_name}_IT Infrastructure, HW, and VM Design {file_extension}'
    bin_str = bin_file.read()
    href = f'<a href="data:application/octet-stream;base64,{base64.b64encode(bin_str).decode()}" download="{file_name}">{file_label}</a>'
    return href

def add_bullet_points_with_bold(doc, text, heading, bullet_symbol="•", force_bullet=False):
    """
    Adds a section with bullet points and handles bold text within the points.
    Ensures consistent bullet style using a specified bullet symbol.
    Removes leading symbols (like '-') from each line.
    Skips bullets for lines that are entirely bold or start with bold text unless forced.
    Skips adding the heading if it is empty or contains only whitespace.
    """
    # Add the section heading if it's not empty or whitespace
    if heading.strip():
        doc.add_heading(heading, level=3)

    # Process each line in the text
    for line in text.strip().split('\n'):
        if line.strip():  # Ensure the line isn't empty
            # Remove leading symbols like '-' or bullets
            cleaned_line = line.strip().lstrip('-').strip()

            # Check if the line starts with bold text
            if not force_bullet and cleaned_line.startswith("**") and "**" in cleaned_line[2:]:
                # Add a paragraph with bold formatting but no bullet
                parts = cleaned_line.split("**")
                paragraph = doc.add_paragraph()
                for i, part in enumerate(parts):
                    run = paragraph.add_run(f"{part.strip()} ")
                    if i % 2 == 1:  # Odd indices are bold text
                        run.bold = True
            else:
                # Add a new paragraph for the bullet point
                paragraph = doc.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(6)  # Add space after each bullet point

                # Insert the bullet symbol
                run = paragraph.add_run(f'{bullet_symbol} ')
                run.bold = False  # Bullet itself is not bold

                # Split the cleaned line to handle bold formatting
                parts = cleaned_line.split("**")  # Handle bold segments
                for i, part in enumerate(parts):
                    run = paragraph.add_run(f"{part.strip()} ")
                    if i % 2 == 1:  # Odd indices are bold text
                        run.bold = True


def generate_document_from_template(
    template_path,
    results,
    results_grade1,
    results_grade3,
    df_comparison,
    third_party_licenses,
    notes,
    input_table,
    customer_name,
    high_availability,
    server_specs,
    gpu_specs,
    first_year_storage_raid5,
    total_image_storage_raid5,
    num_studies,
    storage_title,
    shared_storage,
    raid_1_storage_tb,
    gateway_specs,
    workstation_notes=None,
    diagnostic_specs=None,
    review_specs=None,
    clinician_specs=None,
    ris_specs=None,
    project_grade=None,
    storage_table=None,
    physical_design=None,
    nas_backup_details=None,
    tier_2_disks=None,
    tier_2_disk_size=None,
    tier_3_disks=None,
    tier_3_disk_size=None,
    additional_vm_table=None,
    additional_vm_notes=None,
    general_notes=None,
    additional_servers=None,
    additional_vms=None,
    additional_requirements_table=None
):
    """
    Generates a Word document based on the given inputs and template.
    """

    doc = Document(template_path)

    # Add Title

    title_text = f"{customer_name} IT Infrastructure, HW, and VM Design Recommendations for Medical Imaging Solutions"
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add introduction with bold formatting
    intro_text = generate_introduction(
        customer_name, high_availability, nas_backup_details, tier_3_disks, gpu_specs)

    # Add Introduction to Document
    add_introduction_to_document(doc, intro_text)

    description_text = (
        "This section outlines the **client requirements** and **system assumptions** that serve as the foundation "
        "for the proposed design. Key inputs such as study sizes, annual growth rates, and system capacity are "
        "explicitly detailed to ensure alignment with the operational goals and scalability requirements of the client. "
        "The design is crafted to address both current needs and future expansions, ensuring a robust and adaptable "
        "infrastructure tailored to PaxeraHealth imaging solutions."
    )

    add_table_with_no_split(
        doc,
        input_table,
        title='Client Requirements and System Assumptions',
        description=description_text
    )

    # Utility to add tables
    if storage_table is not None:
        # Base description
        storage_description = (
            "The table below outlines the **Storage Architecture and Scalability Plan**, designed to meet the operational "
            "and performance requirements of the system. The storage solution is categorized into distinct tiers to ensure "
            "optimized data access and long-term reliability:\n\n"
        )

        # Tier 1 Description
        tier_1_description = (
            "- **Tier 1: RAID 1 SSD (NVMe or M.2)**: High-speed, fault-tolerant storage dedicated to the operating system and database. "
            "This tier ensures rapid access to critical system data and overall system stability.\n"
        )

        # Check for Tier 2: Fast Image Storage in the storage table
        tier_2_short_term_exists = any("Tier 2: Fast Image Storage (SSD RAID 5)" in row for row in storage_table)

        if tier_2_short_term_exists:
            # Tier 2: Short-Term Storage
            tier_2_description = (
                "- **Tier 2: Fast Image Storage (RAID 5 SSD - SAS or NLSAS SSD)**: Optimized for fast retrieval of frequently accessed imaging data. "
                "This tier is configured with high-performance SSDs to ensure rapid data access for recent studies, typically retained for 6 months or 1 year.\n"
            )

            # Tier 3: Long-Term Storage
            tier_3_description = (
                "- **Tier 3: Long-Term Storage (RAID 5 or 6 - SATA or NLSAS HDD)**: Scalable storage optimized for archiving historical imaging data. "
                "This tier ensures data protection through redundancy and supports seamless access for low-frequency retrieval needs.\n\n"
            )
        else:
            # If no short-term storage, rename long-term storage as Tier 2
            tier_2_description = (
                "- **Tier 2: Long-Term Storage (RAID 5 or 6 - SATA or NLSAS HDD)**: Scalable storage optimized for archiving historical imaging data. "
                "This tier ensures data protection through redundancy and supports seamless access for low-frequency retrieval needs.\n\n"
            )
            tier_3_description = ""  # No Tier 3 if no short-term storage

        # Compile descriptions
        storage_description += tier_1_description
        storage_description += tier_2_description
        if tier_2_short_term_exists:
            storage_description += tier_3_description

        import pandas as pd

        # Check if Tier 3 column exists and contains only NaN values
        if "Tier 3: Long-Term Storage (HDD RAID 5)" in storage_table.columns:
            if storage_table["Tier 3: Long-Term Storage (HDD RAID 5)"].isna().all():
                # Drop Tier 3 column if all values are NaN
                storage_table = storage_table.drop(columns=["Tier 3: Long-Term Storage (HDD RAID 5)"])

        # Add storage description
        storage_description += (
            "This architecture balances performance, scalability, and cost-efficiency, ensuring seamless availability of imaging data to "
            "support both operational and long-term requirements."
        )

        # Adding the updated table and description to the document
        add_table_with_no_split(
            doc,
            storage_table,
            title='Storage Architecture and Scalability Plan',
            description=storage_description
        )

    # Add VM Recommendations
    if not results.empty:
        # Modify the DataFrame to adjust total row content
        results = results.iloc[:-2]
        results.at[results.index[-1], "Operating System"] = ""
        results.at[results.index[-1], results.columns[0]] = "Total"

        # Core System VMs Section
        gpu_text = "including **AI capabilities**." if gpu_specs else ""
        vm_description_text = (
            f"The following table specifies the **Core Virtual Machine (VM) Configurations** required for **production workloads**. "
            f"These configurations are designed to ensure **high performance**, **scalability**, and **reliability** for the core system {gpu_text}"
        )
        add_table_with_no_split(
            doc,
            results,
            title='Core Virtual Machine Configurations',
            description=vm_description_text
        )

        # General Notes Section (formerly vCore Notes)
        general_notes = [
            "Each **vCore** represents a thread in the physical processor.",
            "Physical processors typically support **two threads per core**, resulting in a **1:2 ratio** of physical cores to vCores.",
            "This configuration ensures efficient utilization of processor resources, enabling **high performance** and **optimal workload distribution** in virtualized environments.",
        ]
        if gpu_specs:
            # List of AI VM names to check in the results
            ai_vm_names = [
                "Organ Segmentator Docker",
                "Lesion Segmentator 2D Docker",
                "Lesion Segmentator 3D Docker",
                "Speech-to-text Docker",
                "AI ARK Manager",
                "AI ARK LAB"
            ]

            # Extract AI VM names present in the results
            ai_vms_present = [vm for vm in ai_vm_names if vm in results["VM Type"].values]

            # Construct the note dynamically with highlighted AI VMs
            if ai_vms_present:
                general_notes.append(
                    "The **Docker-based AI VMs** listed below are specifically designed for **AI processing tasks**, leveraging GPU acceleration for advanced functionalities:"
                )
                for ai_vm in ai_vms_present:
                    general_notes.append(f"- {ai_vm}")

        doc.add_heading('General Notes', level=3)
        for note in general_notes:
            add_bullet_points_with_bold(doc, note.strip(), "",force_bullet=True)

        # Recommended Resources Section for Core VMs
        recommended_resources_description = (
            "The table below outlines the **recommended system specifications** for the core virtual machines. "
            "These specifications are designed to ensure **reliable** and **efficient operation** of the production environment, "
            "meeting **performance demands** and **scalability requirements**."
        )
        add_table_with_no_split(
            doc,
            df_comparison,
            title='Recommended Resources for Core VMs',
            description=recommended_resources_description.strip()
        )

        # Auxiliary VMs Section (if applicable)
        if additional_vm_table is not None:
            auxiliary_vm_purpose = "management operations only" if not additional_vms else "management and testing/training operations"
            additional_vm_description = (
                f"The table below outlines the **Auxiliary Virtual Machine (VM) Configurations** tailored to handle **{auxiliary_vm_purpose}**. "
                f"These VMs are **isolated** from the production environment to ensure **uninterrupted performance** and **flexibility**."
            )
            add_table_with_no_split(
                doc,
                additional_vm_table,
                title='Auxiliary Virtual Machine Configurations',
                description=additional_vm_description
            )

        # Additional VM Notes
        if additional_vms:
            doc.add_heading('Additional VM Notes', level=3)
            if isinstance(additional_vm_notes, list):
                for note in additional_vm_notes:
                    add_bullet_points_with_bold(doc, note.strip(), "")
            elif isinstance(additional_vm_notes, dict):
                for note in additional_vm_notes.values():
                    add_bullet_points_with_bold(doc, note.strip(), "")

        # Recommended Resources for Auxiliary VMs Section
        if additional_requirements_table is not None:
            additional_vm_description = (
                "The table below outlines the **recommended system specifications** for the auxiliary virtual machines. "
                "These specifications ensure that sufficient resources are allocated to auxiliary tasks without impacting core production workloads."
            )
            add_table_with_no_split(
                doc,
                additional_requirements_table,
                title='Recommended Resources for Auxiliary VMs',
                description=additional_vm_description.strip()
            )

    # Add Physical System Design
    if physical_design:
        # Add main heading for Physical System Design
        doc.add_heading('Physical System Design', level=3)

        # Determine the server setup and structure the introduction
        if high_availability:
            server_setup_text = (
                "- **Server Configuration**: The system is designed with **multiple servers** in a **High Availability (HA)** setup. "
                "These servers are connected to a **shared DAS/SAN (Direct-Attached Storage/Storage Area Network)** to ensure fault tolerance, scalability, and continuous operation."
            )
        else:
            server_setup_text = (
                "- **Server Configuration**: The system uses a **standard server** with **built-in storage**, optimized for cost-effectiveness while maintaining reliability."
            )

        # Include GPUs if selected
        gpu_text = "- **GPU-Enabled Servers**: Designed to support advanced AI functionalities such as segmentators and speech-to-text modules." if gpu_specs else ""

        # Include additional VMs dynamically based on selection
        additional_vm_text = ""
        if additional_vms:
            vm_purpose = []
            if any(vm["VM Type"] == "Test Environment VM (Ultima, PACS, Broker)" for vm in additional_vms):
                vm_purpose.append("testing and training")
            if any(vm["VM Type"] == "Management VM (Backup, Antivirus, vCenter)" for vm in additional_vms):
                vm_purpose.append("management")
            additional_vm_text = f"- **Additional Servers**: Configured to support {', '.join(vm_purpose)} operations. These servers ensure flexibility and uninterrupted production workloads."

        # AI Modules based on selection
        selected_ai_modules = [
            "Organ Segmentator Docker",
            "Lesion Segmentator 2D Docker",
            "Lesion Segmentator 3D Docker",
            "Speech-to-text Docker",
            "AI ARK Manager",
            "AI ARK LAB"
        ]
        ai_modules_in_use = [module for module in selected_ai_modules if module in physical_design]
        ai_text = (
            "- **AI Modules**: The system includes the following AI components: "
            + ", ".join(f"**{module}**" for module in ai_modules_in_use)
            if ai_modules_in_use
            else ""
        )

        # Include NAS backup solution if selected
        nas_text = "- **Backup Solution**: Includes a **NAS backup system** to provide redundancy and long-term data protection." if nas_backup_details else ""

        # Combine all system components into structured bullet points
        system_components = [
            server_setup_text,
            gpu_text,
            additional_vm_text,
            ai_text,
            nas_text,
        ]
        system_components = [comp for comp in system_components if comp]  # Filter out empty lines

        # Add structured introduction for system setup
        add_paragraph_with_bold(doc, "The proposed system comprises the following hardware setup:")
        for component in system_components:
            add_bullet_points_with_bold(doc, component.strip(), "", force_bullet=True)

        # Add adaptability note
        adaptability_note = (
            "This hardware setup can be deployed in the client's data center, provided equivalent resources and functionalities are allocated "
            "to maintain similar performance and operational reliability."
        )
        add_paragraph_with_bold(doc, adaptability_note)

        # Process Physical Design Details
        if physical_design.strip():
            design_lines = physical_design.strip().split("\n")
            if design_lines:
                subheading = design_lines[0].strip().rstrip(":")
                remaining_lines = "\n".join(design_lines[1:])
                add_bullet_points_with_bold(doc, remaining_lines, subheading)

        # Add Storage Design

        # Add Storage Design
        doc.add_heading('Storage Design', level=3)

        # Storage Type
        storage_design_text = (
            "The storage system leverages **DAS/SAN (Direct-Attached Storage/Storage Area Network)** configurations to support high availability and scalability. "
            "This design is tailored to accommodate large-scale medical imaging workloads while ensuring redundancy and continuous access to critical data."
        ) if high_availability else (
            "The storage system uses **built-in storage** with RAID configurations to provide reliable and cost-effective solutions for smaller-scale operations. "
            "This design focuses on meeting the operational needs of single-server environments."
        )
        add_paragraph_with_bold(doc, storage_design_text)

        # Add Storage Tiers
        tier_1_text = f"""
          **Tier 1**: OS & DB (SSD RAID 1)
          SSD Drives: 2x {raid_1_storage_tb:.2f} TB
          """

        if tier_2_disks is not None and tier_2_disk_size is not None and tier_3_disk_size is not 0:
            # Add Tier 2 and Tier 3 as separate sections
            tier_2_text = f"""
              **Tier 2**: Fast Image Storage (SSD RAID 5)
              SSD Drives: {tier_2_disks}x {tier_2_disk_size:.2f} TB
              """
            tier_3_text = f"""
              **Tier 3**: Long-Term Storage (HDD RAID 5)
              HDD Drives: {tier_3_disks}x {tier_3_disk_size:.2f} TB
              """
            add_bullet_points_with_bold(doc, tier_1_text + tier_2_text + tier_3_text, "")
        else:
            # Promote Tier 3 to Tier 2 if Tier 2 doesn't exist
            if tier_3_disks is  0 and tier_3_disk_size is  0:
                tier_2_promoted_text = f"""
                  **Tier 2**: Long-Term Storage (HDD RAID 5)
                  HDD Drives: {tier_2_disks}x {tier_2_disk_size:.2f} TB
                  """
                add_bullet_points_with_bold(doc, tier_1_text + tier_2_promoted_text, "")
            else:
                # Handle cases where no Tier 2 or Tier 3 details are provided
                add_bullet_points_with_bold(doc, tier_1_text, "")

        # Add Test Server Details
        if additional_vms:
            doc.add_heading('Test and Management Servers', level=3)
            for server in additional_servers:
                server_details = f"""
                **Server:**
                - **Processors**: {server['Processors']}
                - **Total CPU**: {server['Total Cores']} Cores / {server['Total Threads']} Threads
                - **RAM**: {server['RAM']} GB
                """
                add_bullet_points_with_bold(doc, server_details.strip(), "", force_bullet=True)
        # Add NAS Backup Details
        if nas_backup_details:
            doc.add_heading('Backup Storage (NAS)', level=3)
            nas_backup_text = (
                "The NAS backup solution is designed to provide additional redundancy and facilitate seamless data recovery. "
                "This backup configuration ensures long-term protection for critical medical imaging data, safeguarding against "
                "unexpected system failures or data loss."
            )
            add_paragraph_with_bold(doc, nas_backup_text)

            # Split the backup details into lines and pass all lines to bullet points
            backup_lines = nas_backup_details.strip().split("\n")

            if backup_lines:
                add_bullet_points_with_bold(doc, "\n".join(backup_lines), "",
                                            force_bullet=True)  # Add all lines as bullets

            # Add GPU Requirements (if provided)
            # GPU Requirements Section
            if gpu_specs:
                doc.add_heading('GPU Requirements', level=3)

                # Add descriptive text for GPU Requirements
                add_paragraph_with_bold(
                    doc,
                    "The system leverages powerful GPUs to support advanced AI functionalities, enabling seamless processing "
                    "of resource-intensive tasks like segmentation and speech-to-text analysis. These GPUs are optimized for "
                    "high-performance medical imaging workloads."
                )

                # Add GPU details as bullet points with bold formatting
                add_bullet_points_with_bold(doc, gpu_specs, "", force_bullet=True)
    description_text = (
                    "The table below lists the third-party licenses required to support the proposed infrastructure, "
                    "ensuring compatibility and reliability."
     )
    add_table_with_no_split(doc, third_party_licenses, title='Third Party Licenses', description=description_text)
    # Add Notes Section
    add_bullet_points_with_bold(doc, notes['licensing_notes'], 'Licensing Notes',force_bullet=True)
    add_bullet_points_with_bold(doc, notes['sizing_notes'], 'Sizing Notes')
    add_bullet_points_with_bold(doc, notes['technical_requirements'], 'Technical Requirements')
    add_bullet_points_with_bold(doc, notes['network_requirements'], 'Network Requirements (LAN)')


    # Add Gateway Specifications (if provided)
    if gateway_specs:
        doc.add_heading('Gateway Specifications', level=3)

        # Add a descriptive paragraph for the Gateway section
        gateway_text = (
            "The gateway is responsible for acquiring and transmitting images and data from the modalities "
            "to the PACS system at the main site. This ensures secure, efficient data transfer and seamless system interoperability."
        )
        add_paragraph_with_bold(doc, gateway_text)

        # Add Gateway details as bullet points using the existing utility function
        add_bullet_points_with_bold(doc, gateway_specs, "")
    # Workstation Specifications Section
    if any(spec is not None for spec in [diagnostic_specs, ris_specs, review_specs, clinician_specs,workstation_notes]):
        doc.add_heading('Workstation Specifications', level=2)

        # Add descriptive text for workstation specifications
        add_paragraph_with_bold(
            doc,
            "The proposed workstations are tailored to meet the demands of high-resolution medical imaging and efficient "
            "clinical workflows. Each workstation is equipped with state-of-the-art hardware to ensure optimal performance "
            "and user satisfaction."
        )

        # Add Diagnostic Workstation Specifications
        if diagnostic_specs is not None:
            add_table_with_no_split(doc, diagnostic_specs, title='Diagnostic Workstation')

        # Add RIS Workstation Specifications
        if ris_specs is not None:
            add_paragraph_with_bold(doc, "**RIS Workstation Specifications**:")
            add_table_with_no_split(doc, ris_specs, title='RIS Workstation')

        # Add Review Workstation Specifications
        if review_specs is not None:
            add_paragraph_with_bold(doc, "**Review Workstation Specifications**:")
            add_table_with_no_split(doc, review_specs, title='Review Workstation')

        # Add Clinician Workstation Specifications
        if clinician_specs is not None:
            add_paragraph_with_bold(doc, "**Clinician Workstation Specifications**:")
            add_table_with_no_split(doc, clinician_specs, title='Clinician Workstation')
        add_bullet_points_with_bold(doc, workstation_notes, ' Workstation Requirements & Recommendations')

    add_bullet_points_with_bold(doc, notes['minimum_requirements'], ' Requirements & Recommendations')

    # Save to a buffer and return
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
