import os
import sys
from io import BytesIO
import streamlit as st
import pandas as pd
import numpy as np
from PIL import Image  # Import the Image module from Pillow
import base64

from docx import Document

import base64
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_binary_file_downloader_html(bin_file, file_label='File', customer_name='Customer', file_extension='.docx'):
    file_name = f'{customer_name}_vm_results{file_extension}'
    bin_str = bin_file.read()
    href = f'<a href="data:application/octet-stream;base64,{base64.b64encode(bin_str).decode()}" download="{file_name}">{file_label}</a>'
    return href
def generate_document_from_template(template_path, results, results_grade1, results_grade3, df_comparison,
                                    third_party_licenses, notes, input_table, customer_name, high_availability, server_specs, gpu_specs):
    doc = Document(template_path)  # Load the Word template

    # Add a title with Customer Name centered
    title = doc.add_heading(f'{customer_name} HW Recommendation', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Customer Information', level=2)
    doc.add_paragraph(f'Customer Name: {customer_name}')
    doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d")}')

    # Function to add tables to the document with custom styles
    def add_table(df, heading, style_name):
        doc.add_heading(heading, level=2)
        table = doc.add_table(rows=1, cols=len(df.columns), style=style_name)
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    # Modify input_table for customer load or technical inputs
    input_table = input_table[~input_table['Input'].isin(['Project Grade'])]
    input_table['Value'] = input_table['Value'].replace({True: 'Required', False: 'Not Required'})

    # Add modified input values as a table
    add_table(input_table, 'Customer Load / Technical Inputs', 'CustomTableStyle')

    # Adding tables to the document
    add_table(results, 'VM Recommendations', 'CustomTableStyle')
    add_table(df_comparison, 'Minimum vs. Recommended Resources', 'CustomTableStyle')
    add_table(third_party_licenses, 'Third Party Licenses', 'CustomTableStyle')

    # Function to add bullet points
    def add_bullet_points(text, heading):
        doc.add_heading(heading, level=2)
        for line in text.strip().split('\n'):
            if line.strip():
                paragraph = doc.add_paragraph(line.strip())
                p = paragraph._element
                pPr = p.get_or_add_pPr()
                numPr = OxmlElement('w:numPr')
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), '1')
                numPr.append(numId)
                pPr.append(numPr)

    # Add notes as bullet points
    add_bullet_points(notes['sizing_notes'].replace('-', ''), 'Sizing Notes')
    add_bullet_points(notes['technical_requirements'].replace('-', ''), 'Technical Requirements')
    add_bullet_points(notes['network_requirements'].replace('-', ''), 'Network Requirements (LAN)')

    # Add High Availability or Server Design as bullet points
    design_heading = 'High Availability Design' if high_availability else 'Server Design'
    add_bullet_points(server_specs.replace('-', ''), design_heading)

    # Add GPU Requirements if applicable
    if gpu_specs:
        add_bullet_points(gpu_specs.replace('-', ''), 'GPU Requirements')

    # Save the document to a buffer and return it
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
modality_sizes = {
    "CT": 1.0240,
    "MR": 0.100,
    "US": 0.010,
    "NM": 0.010,
    "X-ray": 0.030,
    "MG": 0.16,
    "Cath": 0.300
}
def calculate_referring_physician_resources(ref_phys_ccu):
    """Calculates the Referring Physician VM configuration based on the CCU.

    Args:
        ref_phys_ccu: The number of Referring Physician concurrent users.

    Returns:
        A list of tuples [(num_vms, ram_gb, vcores)], each representing a VM configuration.
    """

    # Referring Physician-specific configuration ranges
    ccu_thresholds = [8, 16, 24,32, 48, 64]  # Doubled CCU thresholds
    ram_gb_tiers = [8, 16, 24, 32,48, 64]   # Same RAM tiers as Ultima/RIS
    vcores_tiers = [4, 6, 8, 10,10, 12]    # Same vCores tiers as Ultima/RIS
    max_ccu_per_vm = 64  # Maximum CCU per VM (doubled from Ultima)

    if ref_phys_ccu <= max_ccu_per_vm:
        # Single VM case: find the appropriate tier
        for i, threshold in enumerate(ccu_thresholds):
            if ref_phys_ccu <= threshold:
                return [(1, ram_gb_tiers[i], vcores_tiers[i])]
    else:
        # Multiple VM case:
        num_full_vms = ref_phys_ccu // max_ccu_per_vm
        remaining_ccu = ref_phys_ccu % max_ccu_per_vm

        vm_configs = []

        # Full VMs with max specs
        for _ in range(num_full_vms):
            vm_configs.append((1, ram_gb_tiers[-1], vcores_tiers[-1]))

        if remaining_ccu > 0:
            # Last VM with remaining CCUs
            for i, threshold in enumerate(ccu_thresholds):
                if remaining_ccu <= threshold:
                    vm_configs.append((1, ram_gb_tiers[i], vcores_tiers[i]))
                    break

        return vm_configs

def calculate_vm_specifications(vm_type, num_vms, pacs_ccu, ris_ccu, ref_phys_ccu, project_grade):
    vm_specs = {
        "PaxeraUltima": {"vcores": 8, "base_ram": 8, "storage_gb": 150},
        "PaxeraPACS": {"vcores": 8, "base_ram": 16, "storage_gb": 150},
        "PaxeraBroker": {"vcores": 8, "base_ram": 16, "storage_gb": 150},
        "PaxeraRIS": {"vcores": 8, "base_ram": 8, "storage_gb": 150},
        "DBServer": {"vcores": 12, "base_ram": 32, "storage_gb": 400},
        "Referring Physician": {"vcores": 8, "base_ram": 8, "storage_gb": 150},
    }

    vm_requirements = {}
    max_ccu_per_vm = 32

    def get_vm_specs(ccu, base_ram, base_vcores):
        if ccu <= 2:
            ram = 8
            vcores = 6
        else:
            ram = min(64, base_ram + ((ccu - 2) * 2))
            vcores = min(12, base_vcores + ((ccu - 2) // 2) * 2)
        return vcores, ram

    if vm_type in ["PaxeraUltima", "PaxeraRIS"]:
        ccu = pacs_ccu if vm_type == "PaxeraUltima" else ris_ccu
        vm_count = (ccu + max_ccu_per_vm - 1) // max_ccu_per_vm
        for i in range(vm_count):
            vm_name = f"{vm_type}0{i + 1}"
            ccu_for_vm = min(max_ccu_per_vm, ccu - i * max_ccu_per_vm)
            vcores, ram = get_vm_specs(ccu_for_vm, vm_specs[vm_type]["base_ram"], vm_specs[vm_type]["vcores"])
            vm_requirements[vm_name] = {
                "VM Type": vm_type,
                "vCores": vcores,
                "RAM (GB)": ram,
                "Storage (GB)": vm_specs[vm_type]["storage_gb"],
            }
    else:
        for i in range(num_vms):
            vm_name = f"{vm_type}0{i + 1}"
            base_ram = vm_specs[vm_type].get("base_ram", 32)
            vm_requirements[vm_name] = {
                "VM Type": vm_type,
                "vCores": vm_specs[vm_type]["vcores"],
                "RAM (GB)": base_ram,
                "Storage (GB)": vm_specs[vm_type]["storage_gb"],
            }

    return vm_requirements


def calculate_dbserver_resources(num_studies):
    """Calculates the Database VM configuration based on the number of studies,
    ensuring no duplicate configurations.
    """

    max_ram_gb = 64
    max_vcores = 12
    max_studies_single_vm = 300000
    db_tiers = {
        5000: (8, 16),
        50000: (10, 32),
        300000: (12, 64)
    }

    if num_studies <= max_studies_single_vm:
        # Single VM case: linear scaling
        prev_studies_threshold = 0
        prev_base_vcores, prev_base_ram_gb = 0, 0
        for studies_threshold, (base_vcores, base_ram_gb) in db_tiers.items():
            if num_studies <= studies_threshold:
                if studies_threshold != 5000:
                    scaling_factor = (num_studies - prev_studies_threshold) / (studies_threshold - prev_studies_threshold)
                else:
                    scaling_factor = num_studies / studies_threshold
                ram_gb = prev_base_ram_gb + (base_ram_gb - prev_base_ram_gb) * scaling_factor
                vcores = prev_base_vcores + (base_vcores - prev_base_vcores) * scaling_factor
                return [(1, int(round(ram_gb)), int(round(vcores)))]
            prev_studies_threshold, prev_base_vcores, prev_base_ram_gb = studies_threshold, base_vcores, base_ram_gb

    else:
        # Multiple VM case: doubling VMs when needed (no duplicates)
        vm_configs = []
        remaining_studies = num_studies
        while remaining_studies > 0:
            vm_studies = min(max_studies_single_vm, remaining_studies)
            remaining_studies -= vm_studies

            # Find the appropriate tier for this VM
            for studies_threshold, (base_vcores, base_ram_gb) in db_tiers.items():
                if vm_studies <= studies_threshold:
                    # Add only if configuration is not a duplicate
                    config = (1, base_ram_gb, base_vcores)
                    if config not in vm_configs:
                        vm_configs.append(config)
                    break

        return vm_configs
def calculate_paxera_ultima_resources(pacs_ccu):
    """Calculates the PaxeraUltima VM configuration based on the PACS CCU.

    Args:
        pacs_ccu: The number of PACS concurrent users.

    Returns:
        A list of tuples [(num_vms, ram_gb, vcores)], each representing a VM configuration.
    """

    # Ultima-specific configuration ranges (adjusted)
    ccu_thresholds = [4, 8, 12, 24, 32]
    ram_gb_tiers = [8, 16, 24, 48, 64]
    vcores_tiers = [4, 6, 8, 10, 12]
    max_ccu_per_vm = 32  # Maximum CCU per VM

    if pacs_ccu <= max_ccu_per_vm:
        # Single VM case: find the appropriate tier
        for i, threshold in enumerate(ccu_thresholds):
            if pacs_ccu <= threshold:
                return [(1, ram_gb_tiers[i], vcores_tiers[i])]
    else:
        # Multiple VM case:
        num_full_vms = pacs_ccu // max_ccu_per_vm
        remaining_ccu = pacs_ccu % max_ccu_per_vm

        vm_configs = []

        # Full VMs with max specs
        for _ in range(num_full_vms):
            vm_configs.append((1, ram_gb_tiers[-1], vcores_tiers[-1]))

        if remaining_ccu > 0:
            # Last VM with remaining CCUs
            for i, threshold in enumerate(ccu_thresholds):
                if remaining_ccu <= threshold:
                    vm_configs.append((1, ram_gb_tiers[i], vcores_tiers[i]))
                    break

        return vm_configs

def calculate_paxera_pacs_resources(num_studies):
    """Calculates the PaxeraPACS VM configuration based on the number of studies."""

    min_ram_gb = 14
    max_ram_gb = 32
    min_vcores = 8
    max_vcores = 12
    max_studies_per_vm = 50000

    if num_studies <= max_studies_per_vm:
        # Single VM case: linear scaling
        scaling_factor = num_studies / max_studies_per_vm
        ram_gb = min_ram_gb + (max_ram_gb - min_ram_gb) * scaling_factor
        vcores = min_vcores + (max_vcores - min_vcores) * scaling_factor
        return [(1, int(round(ram_gb)), int(round(vcores)))]
    else:
        # Multiple VM case:
        num_full_vms = num_studies // max_studies_per_vm
        remaining_studies = num_studies % max_studies_per_vm

        vm_configs = []

        # Full VMs with max specs
        for _ in range(num_full_vms):
            vm_configs.append((1, max_ram_gb, max_vcores))

        if remaining_studies > 0:
            # Last VM with remaining studies
            scaling_factor = remaining_studies / max_studies_per_vm
            ram_gb = min_ram_gb + (max_ram_gb - min_ram_gb) * scaling_factor
            vcores = min_vcores + (max_vcores - min_vcores) * scaling_factor
            vm_configs.append((1, int(round(ram_gb)), int(round(vcores))))

        return vm_configs
def calculate_vm_requirements(num_studies, pacs_ccu, ris_ccu, ref_phys_ccu, project_grade, broker_required,
                              contract_duration, study_size_mb, annual_growth_rate, breakdown_per_modality=False,
                              aidocker_included=False, ark_included=False,
                              **modality_cases):
    vm_specs = {
        "PaxeraUltima": {"vcores": 8, "base_ram": 8, "storage_gb": 150},
        "PaxeraPACS": {"vcores": 8, "base_ram": 16, "storage_gb": 150},
        "PaxeraBroker": {"vcores": 8, "base_ram": 16, "storage_gb": 150},
        "PaxeraRIS": {"vcores": 8, "base_ram": 8, "storage_gb": 150},
        "DBServer": {"vcores": 12, "base_ram": 32, "storage_gb": 400},
        "Referring Physician": {"vcores": 8, "base_ram": 8, "storage_gb": 150},
    }

    vms_needed = {
        "PaxeraUltima": 1,
        "DBServer": 1,
        "PaxeraBroker": 1 if broker_required or (pacs_ccu > 0 and ris_ccu > 0) else 0,
        "PaxeraRIS": 1 if ris_ccu > 0 else 0,
        "Referring Physician": 1 if ref_phys_ccu > 0 else 0,
    }

    if breakdown_per_modality:
        total_cases = sum(modality_cases.values())
        total_storage = sum(modality_cases[modality] * modality_sizes.get(modality, 0) for modality in modality_cases)

        # Calculate average study size from breakdown
        average_study_size = total_storage / total_cases if total_cases > 0 else 0

        # Calculate RAID 5 storage based on modality breakdown
        image_storage_raid5_modality = sum(
            modality_cases[modality] * modality_sizes.get(modality, 0) * contract_duration for modality in
            modality_cases
        )
    else:
        total_cases = num_studies
        average_study_size = study_size_mb

        # Calculate RAID 5 storage based on total studies
        image_storage_raid5_modality = num_studies * average_study_size * contract_duration * (
                1 + annual_growth_rate / 100)

    vm_requirements = {}
    vm_config_lists = {
        "PaxeraUltima": calculate_paxera_ultima_resources(pacs_ccu),
        "PaxeraRIS": calculate_paxera_ultima_resources(ris_ccu) if ris_ccu > 0 else [],
        "Referring Physician": calculate_referring_physician_resources(ref_phys_ccu) if ref_phys_ccu > 0 else []
    }

    for vm_type, config_list in vm_config_lists.items():
        for i, (num_vms, ram_gb, vcores) in enumerate(config_list):
            vm_name = f"{vm_type}{i + 1:02d}"
            vm_requirements[vm_name] = {
                "VM Type": vm_type,
                "vCores": vcores,
                "RAM (GB)": ram_gb,
                "Storage (GB)": 150,  # Adjust if needed for different VM types
            }

    # Assign DBServer specs to match PaxeraPACS
    dbserver_vm_configs = calculate_dbserver_resources(num_studies)
    # Assuming only one DB server is needed
    _, ram_gb, vcores = dbserver_vm_configs[0]
    vm_requirements["DBServer01"] = {
        "VM Type": "DBServer",
        "vCores": vcores,
        "RAM (GB)": ram_gb,
        "Storage (GB)": 400  # Assuming 400GB storage for DBServer
    }
    # Calculate PaxeraPACS specs using the new function
    pacs_vm_configs = calculate_paxera_pacs_resources(num_studies)

    # Update the vm_requirements with the PaxeraPACS VM specs
    for i, (num_vms, ram_gb, vcores) in enumerate(pacs_vm_configs):
        vm_name = f"PaxeraPACS{i + 1:02d}"
        vm_requirements[vm_name] = {
            "VM Type": "PaxeraPACS",
            "vCores": vcores,
            "RAM (GB)": ram_gb,
            "Storage (GB)": 150,  # Assuming 150GB storage for each PaxeraPACS VM
        }
    # Calculate other VM specifications (EXCLUDING PaxeraUltima)
    for vm_type in ["PaxeraBroker"]:
        if vms_needed.get(vm_type, 0) > 0:
            vm_specs_combined = calculate_vm_specifications(vm_type, vms_needed[vm_type], pacs_ccu, ris_ccu,
                                                            ref_phys_ccu, project_grade)
            vm_requirements.update(vm_specs_combined)

    # Project Grade 1: Combine PaxeraPACS and PaxeraUltima
    if project_grade == 1:
        if "PaxeraPACS01" in vm_requirements and "PaxeraUltima01" in vm_requirements:
            combined_vm = {
                "VM Type": "PaxeraPACS/PaxeraUltima",
                "vCores": 2 * round((vm_requirements["PaxeraPACS01"]["vCores"] + vm_requirements["PaxeraUltima01"][
                    "vCores"]) / 1.5 / 2),
                "RAM (GB)": 2 * round((vm_requirements["PaxeraPACS01"]["RAM (GB)"] + vm_requirements["PaxeraUltima01"][
                    "RAM (GB)"]) / 1.5 / 2),
                "Storage (GB)": vm_requirements["PaxeraPACS01"]["Storage (GB)"] + vm_requirements["PaxeraUltima01"][
                    "Storage (GB)"]
            }
            vm_requirements["PaxeraPACS01"] = combined_vm
            del vm_requirements["PaxeraUltima01"]

        # Combine DBServer and PaxeraBroker if necessary
        if broker_required or (pacs_ccu > 0 and ris_ccu > 0):
            if "DBServer01" in vm_requirements and "PaxeraBroker01" in vm_requirements:
                combined_vm = {
                    "VM Type": "DBServer/PaxeraBroker",
                    "vCores": min(12, 2 * round((vm_requirements["DBServer01"]["vCores"] +
                                                 vm_requirements["PaxeraBroker01"]["vCores"]) / 1.5 / 2)),
                    "RAM (GB)": min(64, 2 * round((vm_requirements["DBServer01"]["RAM (GB)"] +
                                                   vm_requirements["PaxeraBroker01"]["RAM (GB)"]) / 1.5 / 2)),
                    "Storage (GB)": vm_requirements["DBServer01"]["Storage (GB)"] + vm_requirements["PaxeraBroker01"][
                        "Storage (GB)"]
                }
                vm_requirements["DBServer01"] = combined_vm
                del vm_requirements["PaxeraBroker01"]

    elif project_grade == 2:
        # Project Grade 2 adjustments (combine DBServer and PaxeraBroker ONLY)
        if "DBServer01" in vm_requirements and "PaxeraBroker01" in vm_requirements:
            combined_vm = {
                "VM Type": "DBServer/PaxeraBroker",
                "vCores": min(12, 2 * round((vm_requirements["DBServer01"]["vCores"] +
                                             vm_requirements["PaxeraBroker01"]["vCores"]) / 1.5 / 2)),
                "RAM (GB)": min(64, 2 * round((vm_requirements["DBServer01"]["RAM (GB)"] +
                                               vm_requirements["PaxeraBroker01"]["RAM (GB)"]) / 1.5 / 2)),
                "Storage (GB)": vm_requirements["DBServer01"]["Storage (GB)"] +
                                vm_requirements["PaxeraBroker01"]["Storage (GB)"]
            }
            vm_requirements["DBServer01"] = combined_vm
            del vm_requirements["PaxeraBroker01"]
    # Project Grade 3: Increase resources for each VM except PaxeraPACS
    elif project_grade == 3:
        # Project Grade 3 adjustments: NO COMBINATIONS
        for vm_name, specs in vm_requirements.items():
            if specs["VM Type"] != "PaxeraPACS":
                specs["vCores"] = 2 * round(specs["vCores"] * 1 / 2)
                specs["RAM (GB)"] = 2 * round(specs["RAM (GB)"] * 1 / 2)

        # Ensure DBServer and PaxeraBroker are separate, even if broker is required
        if "DBServer/PaxeraBroker" in vm_requirements:
            # Split combined VM back into DBServer and PaxeraBroker
            dbserver_specs = vm_requirements["DBServer/PaxeraBroker"].copy()
            dbserver_specs["VM Type"] = "DBServer"
            vm_requirements["DBServer01"] = dbserver_specs

            broker_specs = vm_requirements["DBServer/PaxeraBroker"].copy()
            broker_specs["VM Type"] = "PaxeraBroker"
            vm_requirements["PaxeraBroker01"] = broker_specs

            del vm_requirements["DBServer/PaxeraBroker"]  # Remove combined VM

    else:
        raise ValueError(
            "Invalid project grade. Please choose 1, 2, or 3.")  # Combine DBServer and PaxeraBroker if necessary

    df_results = pd.DataFrame()
    total_image_storage_raid5 = 0
    for year in range(contract_duration):
        image_storage_raid5 = (num_studies * study_size_mb * (1 + annual_growth_rate / 100)) / 1024
        total_image_storage_raid5 += image_storage_raid5
        num_studies *= (1 + annual_growth_rate / 100)  # Increment num_studies for the next year
    total_image_storage_raid5 = round(total_image_storage_raid5, 2)
    total_vcpu = sum([vm_requirements[vm_name]["vCores"] for vm_name in vm_requirements])
    total_ram = sum([vm_requirements[vm_name]["RAM (GB)"] for vm_name in vm_requirements])
    total_storage = sum([vm_requirements[vm_name]["Storage (GB)"] for vm_name in vm_requirements])

    if vm_requirements:
        df_results = pd.DataFrame.from_dict(vm_requirements, orient='index')
        if aidocker_included:
            vm_requirements["AISegmentationDocker01"] = {
                "VM Type": "AI Segmentation Docker",
                "vCores": 12,
                "RAM (GB)": 32,
                "Storage (GB)": 300
            }
            total_vcpu += 12
            total_ram += 32
            total_storage += 300
        if ark_included:
            # Add two AI Segmentation Docker VMs
            for i in range(2):
                vm_name = f"AISegmentationDocker0{i + 1}"
                vm_requirements[vm_name] = {
                    "VM Type": "AI Segmentation Docker",
                    "vCores": 12,
                    "RAM (GB)": 32,
                    "Storage (GB)": 300
                }

            # Add AI ARK LAB VM
            vm_requirements["AIARKLAB01"] = {
                "VM Type": "AI ARK LAB",
                "vCores": 12,
                "RAM (GB)": 32,
                "Storage (GB)": 300
            }

            # Add resources to totals
            total_vcpu += 36  # 12 vCores * 3 VMs
            total_ram += 96  # 32 GB RAM * 3 VMs
            total_storage += 900  # 300 GB Storage * 3 VMs        # --- Update DataFrame ---
        df_results = pd.DataFrame.from_dict(vm_requirements, orient='index')

        df_results.loc["Total"] = ["-", total_vcpu, total_ram, total_storage]
        df_results.loc["RAID 1 (SSD)"] = ["-", "-", "-", total_storage]
        df_results.loc["RAID 5 (HDD)"] = ["-", "-", "-", round(image_storage_raid5_modality,
                                                               2) if breakdown_per_modality else total_image_storage_raid5]

    if num_studies <= 50000:
        sql_license = "SQL Express"
    elif num_studies <= 200000:
        sql_license = "SQL Standard"
    else:
        sql_license = "SQL Enterprise"

    return df_results, sql_license, image_storage_raid5_modality, total_vcpu, total_ram, total_image_storage_raid5


def main():
    logo_image = Image.open("D:/Presales Documents/logo.png")  # Replace with the actual image file path

    # Display logo using columns for alignment
    col1, col2, col3 = st.columns([1, 6, 1])  # Adjust column ratios for centering
    with col1:
        st.write("")  # Empty column for spacing
    with col2:
        st.image(logo_image)  # Display the logo
    with col3:
        st.write("")  # Empty column for spacing

    st.title("PaxeraHealth VM Calculator")
    customer_name = st.text_input("Customer Name:")

    st.subheader("Input Method:")
    breakdown_per_modality = st.radio("Breakdown per Modality?", ["No", "Yes"])

    if breakdown_per_modality == "No":
        num_studies = st.number_input("Number of studies per year:", min_value=0, value=100000)
        modality_cases = {}
    else:
        st.subheader("Modality Breakdown:")
        modality_cases = {
            "CT": st.number_input("CT Cases:", min_value=0),
            "MR": st.number_input("MR Cases:", min_value=0),
            "US": st.number_input("US Cases:", min_value=0),
            "NM": st.number_input("NM Cases:", min_value=0),
            "X-ray": st.number_input("X-ray Cases:", min_value=0),
            "MG": st.number_input("MG Cases:", min_value=0),
            "Cath": st.number_input("Cath Cases:", min_value=0),
        }
        num_studies = sum(modality_cases.values())

    pacs_ccu = st.number_input("PACS CCU:", min_value=0, value=8)
    ris_ccu = st.number_input("RIS CCU (enter 0 if no RIS):", min_value=0, value=8)
    ref_phys_ccu = st.number_input("Referring Physician CCU (enter 0 if none):", min_value=0, value=8)
    project_grade = st.selectbox("Project Grade:", [1, 2, 3])
    broker_required = st.checkbox("Broker VM Required (check if explicitly requested)", value=False)
    contract_duration = st.number_input("Contract Duration (years):", min_value=1, value=3)
    study_size_mb = st.number_input("Study Size (MB):", min_value=0, value=120)
    annual_growth_rate = st.number_input("Annual Growth Rate (%):", min_value=0.0, value=10.0, format="%f")
    aidocker_included = st.checkbox("Include U9th Integrated AI modules(Auto Segmentation , Spine Labeling, etc.)",
                                    value=False)
    ark_included = st.checkbox("Include ARKAI", value=False)
    high_availability = st.checkbox("High Availability HW Design Required", value=False)

    calculate = st.button("Calculate")

    if calculate:

        results, sql_license, image_storage_raid5_modality, total_vcpu, total_ram, total_storage = calculate_vm_requirements(
            num_studies, pacs_ccu, ris_ccu, ref_phys_ccu, project_grade, broker_required, contract_duration,
            study_size_mb, annual_growth_rate,   # Pass directly
            aidocker_included=aidocker_included, ark_included=ark_included,
            **modality_cases  # Pass as keyword arguments
        )

        if not results.empty:
            # Add new columns "Operating System" and "Other Software"
            results["Operating System"] = "Windows Server 2019 or Higher"
            results["Other Software"] = ""

            # Update "Other Software" for DBServer and DB/Broker
            for index in results.index:
                if "DBServer" in index:
                    results.at[index, "Other Software"] = sql_license
                if "AISegmentationDocker" in index:  # Check for any AI Segmentation Docker VM
                    results.at[index, "Operating System"] = "Ubuntu 20.4"
                    results.at[
                        index, "Other Software"] = "Nvidia Driver version 450.80.02 or higher\nNvidia driver to support CUDA version 11.4 or higher"
                if "AIARKLAB01" in index:
                    results.at[index, "Operating System"] = "Ubuntu 20.4"
                    results.at[index, "Other Software"] = "RTX 4080 / RTX 4090 Video Cards"

            # Remove "Operating System" for the last three rows
            last_three_indices = results.tail(3).index
            results.loc[last_three_indices, "Operating System"] = ""

            # Display the results table with adjusted column width
            st.subheader("VM Recommendations:")
            st.dataframe(results.style.set_properties(subset=["Operating System"], **{'width': '300px'}))

            if breakdown_per_modality == "Yes":
                st.text(f"RAID 5 Storage (Modality breakdown): {round(image_storage_raid5_modality, 2)} GB")
            results_grade1, _, _, total_vcpu_grade1, total_ram_grade1, total_storage_grade1 = calculate_vm_requirements(
                num_studies, pacs_ccu, ris_ccu, ref_phys_ccu, 1, broker_required, contract_duration,
                study_size_mb, annual_growth_rate, breakdown_per_modality, aidocker_included=aidocker_included,
                ark_included=ark_included, **modality_cases
            )
            # Calculate Grade 3 requirements
            results_grade3, _, _, total_vcpu_grade3, total_ram_grade3, total_storage_grade3 = calculate_vm_requirements(
                num_studies, pacs_ccu, ris_ccu, ref_phys_ccu, 3, broker_required, contract_duration,
                study_size_mb, annual_growth_rate, breakdown_per_modality, aidocker_included=aidocker_included,
                ark_included=ark_included, **modality_cases
            )
            # Summarize total vCores and RAM
            total_vcpu_grade1 = results_grade1.loc["Total", "vCores"]
            total_ram_grade1 = results_grade1.loc["Total", "RAM (GB)"]
            total_vcpu_grade3 = results_grade3.loc["Total", "vCores"]
            total_ram_grade3 = results_grade3.loc["Total", "RAM (GB)"]

            raid_1_storage_tb = results.loc["RAID 1 (SSD)", "Storage (GB)"] / 1024
            raid_5_storage_tb = results.loc["RAID 5 (HDD)", "Storage (GB)"] / 1024

            # -- CREATE COMPARISON TABLE --
            comparison_data = {
                "Specification": ["Total vCores", "Total RAM (GB)", "RAID 1 (SSD) (TB)", "RAID 5 (HDD) (TB)"],
                "Minimum Specs (Grade 1)": [total_vcpu_grade1, total_ram_grade1, round(raid_1_storage_tb, 2),
                                            round(raid_5_storage_tb, 2)],
                "Recommended Specs (Grade 3)": [total_vcpu_grade3, total_ram_grade3, round(raid_1_storage_tb, 2),
                                                round(raid_5_storage_tb, 2)]
            }
            df_comparison = pd.DataFrame(comparison_data)

            # -- DISPLAY EVERYTHING --
            # ... (display results table as before, with 'Operating System' and 'Other Software' columns)

            st.subheader("Minimum vs. Recommended Resources:")  # New table title
            st.dataframe(df_comparison.style.set_properties(subset=["Specification"],
                                                            **{'width': '300px'}))  # Show the comparison table
            # High Availability Design
            if high_availability:
                section_header = "High Availability Design:"
                total_vcpu = round(total_vcpu * 1.5 // 2)
                total_vcpu = round(total_vcpu / 2) * 2
                total_core = total_vcpu // 2
                total_core = round(total_core / 2) * 2

                original_vcpu = results.loc["Total", "vCores"]  # Store the original total vCPU
                total_ram = round(total_ram * 1.5)
                nas_backup_storage = round(1.2 * total_storage, 2)
                server_specs = f"""
                    **Server 1:**
                      - CPU: {total_core} Cores({total_vcpu} threads) 
                      - RAM: {total_ram // 2} GB

                    **Server 2:**
                      - CPU: {total_core}  Cores ({total_vcpu} threads)
                      - RAM: {total_ram // 2} GB

                    **Shared DAS/SAN Storage:**
                      - RAID 1 (SSD): {results.loc["RAID 1 (SSD)", "Storage (GB)"]:.2f} GB
                      - RAID 5 (HDD): {total_storage:.2f} GB

                    **NAS Backup Storage:**
                      - RAID 5 (HDD): {nas_backup_storage:.2f} GB
                    """
            else:
                section_header = "Server Design:"
                total_vcpu = results.loc["Total", "vCores"] // 2  # Divide total vCPU by 2
                total_ram = results.loc["Total", "RAM (GB)"]
                server_specs = f"""

                      - CPU: {total_vcpu} Cores ({results.loc["Total", "vCores"]} threads)
                      - Total RAM: {total_ram} GB
                      - Server Built-in Storage:  RAID 1 (SSD): {results.loc["RAID 1 (SSD)", "Storage (GB)"]:.2f} GB   , RAID 5 (HDD): {total_storage:.2f} GB
                    """

            # Display the appropriate design
            st.subheader(section_header)
            st.markdown(server_specs)

            # Third Party Licenses Table
            windows_count = results["Operating System"].value_counts().get("Windows Server 2019 or Higher", 0)
            third_party_licenses = pd.DataFrame({
                "Item Description": [
                    sql_license,
                    "MS Windows Server Standard 2019 or higher",
                    "Antivirus Server License",
                    "VMware vSphere Essentials KIT",
                    "Backup Software"
                ],
                "Qty": [
                    1,
                    windows_count,
                    windows_count,
                    1,
                    1
                ]
            })

            st.subheader("Third Party Licenses")
            st.dataframe(third_party_licenses.style.set_properties(subset=["Item Description"], **{'width': '300px'}))

            # Add sizing notes, technical requirements, and network requirements
            st.subheader("Sizing Notes:")
            st.markdown("""
            - Provided VM sizing of the Virtual servers will be scaled up horizontally or vertically based on the expected volume of data and number of CCUs.
            - SSD Datastore recommended for all OS Virtual disks of all VMs.
            - SSD recommended for the data drive of the Database Server.
            - It's recommended to have SSD storage for the short Term Storage (STS) that keep last 6 month of data for higher performance in data access.
            """)
            gpu_specs = ""
            if aidocker_included:

                st.subheader("GPU Requirements:")
                gpu_specs= """

                    - GPUs: 1
                    - GPU Memory: 32 GB
                    - Nvidia Tesla V100 or equivalent RTX (Preferred)
                    - Nvidia Driver version 450.80.02 or higher
                    - Nvidia driver to support CUDA version 11.4 or higher
                    """
            elif ark_included:
                st.subheader("GPU Requirements:")
                gpu_specs= """
                - GPUs: 3
                - GPU Memory: 32 GB
                -  2* Nvidia Tesla V100 or equivalent RTX (Preferred)(For Segmentation  Dockers)
                - Nvidia Driver version 450.80.02 or higher
                - Nvidia driver to support CUDA version 11.4 or higher
                - RTX 4080 / RTX 4090 Video Cards ( For ARK LAB)
                    """
            st.markdown(gpu_specs)
            st.subheader("Technical Requirements:")
            st.markdown("""
            - Provide remote access to the VMs (all VMs) for SW installation and configurations.
            - In case not using dongles, a connection from the VMs to (https://paxaeraultima.net:1435) for PaxeraHealth licensing except the database VM.
            """)

            st.subheader("Network Requirements (LAN):")
            st.markdown("""
            - LAN bandwidth to be 1GB dedicated.
            - Paxera PACS VMs, Paxera Ultima Viewer VMs, Paxera Broker Integration VMs must be in same vLAN.
            - 1 Gb/s dedicated bandwidth across the vLAN with the maximum number of two network hops.
            - Latency less than 1ms.
            """)
            notes = {
                "sizing_notes": """
                        - Provided VM sizing of the Virtual servers will be scaled up horizontally or vertically based on the expected volume of data and number of CCUs.
                        - SSD Datastore recommended for all OS Virtual disks of all VMs.
                        - SSD recommended for the data drive of the Database Server.
                        - It's recommended to have SSD storage for the short Term Storage (STS) that keep last 6 month of data for higher performance in data access.
                        """,
                "technical_requirements": """
                        - Provide remote access to the VMs (all VMs) for SW installation and configurations.
                        - In case not using dongles, a connection from the VMs to (https://paxaeraultima.net:1435) for PaxeraHealth licensing except the database VM.
                        """,
                "network_requirements": """
                        - LAN bandwidth to be 1GB dedicated.
                        - Paxera PACS VMs, Paxera Ultima Viewer VMs, Paxera Broker Integration VMs must be in same vLAN.
                        - 1 Gb/s dedicated bandwidth across the vLAN with the maximum number of two network hops.
                        - Latency less than 1ms.
                        """
            }

            if calculate:
                # Calculate results
                # ... (results calculation) ...

                # Store the results in session state so they persist across re-runs
                st.session_state.results = results
                st.session_state.calculated = True  # Mark calculation as done

            # Display results only if calculation is done
            if st.session_state.calculated:
                # ... (display results tables, notes, requirements, etc.) ...
                input_values = pd.DataFrame({
                    "Input": ["PACS CCU", "RIS CCU", "Referring Physician CCU", "Project Grade",
                              "Broker VM Required", "Contract Duration (years)", "Study Size (MB)",
                              "Annual Growth Rate (%)"],
                    "Value": [pacs_ccu, ris_ccu, ref_phys_ccu, project_grade, broker_required,
                              contract_duration, study_size_mb, annual_growth_rate, ]
                })
                # Download Word document link
                if not st.session_state.results.empty:
                    doc = generate_document_from_template(
                        os.path.join("assets", "templates", "Temp.docx"),
                        results,
                        results_grade1,
                        results_grade3,
                        df_comparison,
                        third_party_licenses,
                        notes, input_values,
                        customer_name=customer_name,
                        high_availability=high_availability,
                        server_specs=server_specs,
                        gpu_specs=gpu_specs
                    )


                    download_link = get_binary_file_downloader_html(
                        doc,
                        file_label="Download Word Document",
                        customer_name=customer_name,  # Replace "CustomerName" with the actual customer name
                        file_extension=".docx"
                    )
                    st.markdown(download_link, unsafe_allow_html=True)

if __name__ == "__main__":
    main()